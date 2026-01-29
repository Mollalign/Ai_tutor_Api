"""
DOCX Document Parser

Extracts text from Microsoft Word documents using python-docx.

DOCX Structure:
--------------
DOCX files are ZIP archives containing:
- word/document.xml: Main content
- word/styles.xml: Style definitions
- word/media/: Images and media
- [Content_Types].xml: File manifest

python-docx parses these XMLs and provides a nice API.

What We Extract:
---------------
- Paragraphs with their styles (for heading detection)
- Tables (converted to text)
- Lists (with bullets/numbers)
- Document properties (title, author)

What We Don't Extract:
---------------------
- Images (would need OCR)
- Headers/footers (usually not relevant)
- Comments and tracked changes
- Complex formatting
"""

import io
import logging
from typing import List, Optional

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.ai.parsers.base import (
    DocumentParser,
    ParsedDocument,
    PageContent,
    ParserType,
)

logger = logging.getLogger(__name__)


class DOCXParser(DocumentParser):
    """
    Parser for Microsoft Word documents (.docx).
    
    Extracts text, headings, and tables from Word documents.
    
    Note: DOCX doesn't have explicit page breaks like PDF.
    We treat the entire document as one "page" but preserve
    paragraph and heading structure.
    """
    
    @property
    def supported_types(self) -> List[ParserType]:
        return [ParserType.DOCX]
    
    def parse(
        self,
        content: bytes,
        filename: Optional[str] = None
    ) -> ParsedDocument:
        """
        Parse DOCX and extract text.
        
        Args:
            content: Raw DOCX file bytes
            filename: Original filename for logging
        
        Returns:
            ParsedDocument with extracted content
        """
        filename = filename or "unknown.docx"
        logger.info(f"Parsing DOCX: {filename} ({len(content)} bytes)")
        
        try:
            # Create file-like object from bytes
            docx_file = io.BytesIO(content)
            
            # Parse document
            doc = DocxDocument(docx_file)
            
            # Extract metadata
            metadata = self._extract_metadata(doc, filename)
            
            # Extract content
            text_parts = []
            headings = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                para_text = para.text.strip()
                
                if not para_text:
                    continue
                
                # Check if this is a heading
                if para.style and para.style.name:
                    style_name = para.style.name.lower()
                    if 'heading' in style_name or 'title' in style_name:
                        headings.append(para_text)
                        text_parts.append(f"\n## {para_text}\n")
                    else:
                        text_parts.append(para_text)
                else:
                    text_parts.append(para_text)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables, start=1):
                table_text = self._extract_table(table)
                if table_text:
                    text_parts.append(f"\n[Table {table_idx}]\n{table_text}\n")
            
            # Combine text
            full_text = self._clean_text('\n'.join(text_parts))
            
            # Create single page (DOCX doesn't have page structure)
            page_content = PageContent(
                page_number=1,
                text=full_text,
                headings=headings,
                metadata={
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                }
            )
            
            logger.info(
                f"DOCX parsed successfully: {filename}, "
                f"{len(doc.paragraphs)} paragraphs, {len(full_text)} characters"
            )
            
            return ParsedDocument(
                text=full_text,
                pages=[page_content],
                metadata=metadata,
                success=True
            )
            
        except PackageNotFoundError as e:
            error_msg = f"Invalid DOCX file (not a valid Office document): {e}"
            logger.error(f"DOCX parse error for {filename}: {error_msg}")
            return ParsedDocument.from_error(error_msg)
            
        except Exception as e:
            error_msg = f"Error parsing DOCX: {e}"
            logger.exception(f"DOCX parse error for {filename}")
            return ParsedDocument.from_error(error_msg)
    
    def _extract_metadata(
        self,
        doc: DocxDocument,
        filename: str
    ) -> dict:
        """Extract metadata from DOCX document."""
        metadata = {
            "filename": filename,
            "file_type": "docx",
        }
        
        try:
            core_props = doc.core_properties
            
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.created:
                metadata["created"] = str(core_props.created)
            if core_props.modified:
                metadata["modified"] = str(core_props.modified)
                
        except Exception as e:
            logger.debug(f"Could not extract DOCX metadata: {e}")
        
        # Fallback title
        if "title" not in metadata:
            title = filename.rsplit('.', 1)[0]
            title = title.replace('_', ' ').replace('-', ' ')
            metadata["title"] = title
        
        return metadata
    
    def _extract_table(self, table) -> str:
        """
        Convert a DOCX table to text.
        
        Tables are converted to a simple text format:
        | Cell 1 | Cell 2 | Cell 3 |
        | Data 1 | Data 2 | Data 3 |
        
        Args:
            table: python-docx Table object
        
        Returns:
            Text representation of table
        """
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                cells.append(cell_text)
            
            if any(cells):  # Skip empty rows
                rows.append(' | '.join(cells))
        
        return '\n'.join(rows)